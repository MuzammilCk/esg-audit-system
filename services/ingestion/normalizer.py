from __future__ import annotations

import asyncio
import base64
import csv
import json
import os
import zipfile
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import FastAPI, HTTPException

from services.ingestion.models import (
    IngestedDocument,
    NormalizedDocument,
    NormalizedImage,
    NormalizedStructure,
    NormalizedTable,
)
from services.ingestion.parsers.pdf_parser import parse_pdf
from services.ingestion.parsers.xlsx_parser import parse_xlsx


def _postgres_dsn() -> Optional[str]:
    dsn = os.getenv("NORMALIZER_POSTGRES_DSN")
    return dsn if dsn else None


def _postgres_table() -> str:
    return os.getenv("NORMALIZER_POSTGRES_TABLE", "normalized_documents")


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _b64_decode(data: str) -> bytes:
    try:
        return base64.b64decode(data)
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail={
                "error": "invalid_base64",
                "message": "raw_bytes must be base64-encoded.",
                "exception": str(e),
            },
        ) from e


def _detect_kind(*, ingested: IngestedDocument, raw: bytes) -> str:
    mime = (ingested.mime_type or "").lower().strip()
    original_name = str(ingested.metadata.get("original_name") or "")
    ext = Path(original_name).suffix.lower()

    if "pdf" in mime or raw.startswith(b"%PDF"):
        return "pdf"

    if (
        "wordprocessingml.document" in mime
        or ext == ".docx"
        or (mime == "application/zip" and ext == ".docx")
    ):
        return "docx"

    if (
        "spreadsheetml.sheet" in mime
        or ext == ".xlsx"
        or (mime == "application/zip" and ext == ".xlsx")
    ):
        return "xlsx"

    if mime in {"text/csv", "application/csv"} or ext == ".csv":
        return "csv"

    # Zip-based Office docs without reliable mime/ext: inspect contents.
    if raw.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(BytesIO(raw)) as zf:
                names = set(zf.namelist())
            if "word/document.xml" in names:
                return "docx"
            if "xl/workbook.xml" in names:
                return "xlsx"
        except Exception:
            pass

    # Best-effort CSV detection
    try:
        sample = raw[:2048].decode("utf-8")
        if "," in sample and "\n" in sample:
            return "csv"
    except Exception:
        pass

    return "unknown"


def _unstructured_docx_text(raw_docx: bytes) -> str:
    """Best-effort text extraction for DOCX using unstructured."""

    try:
        from unstructured.partition.docx import partition_docx  # type: ignore
    except Exception:
        return ""

    try:
        elements = partition_docx(file=BytesIO(raw_docx))
    except Exception:
        return ""

    parts: List[str] = []
    for el in elements:
        txt = getattr(el, "text", None)
        if txt:
            parts.append(str(txt))
    return "\n".join(parts).strip()


def _parse_docx(raw_docx: bytes) -> Dict[str, Any]:
    from docx import Document

    doc = Document(BytesIO(raw_docx))

    # Text: prefer unstructured, fallback to python-docx paragraphs.
    text_content = _unstructured_docx_text(raw_docx)
    if not text_content:
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        text_content = "\n".join(paragraphs).strip()

    # Tables
    tables: List[NormalizedTable] = []
    for t in doc.tables:
        data: List[List[str]] = []
        for row in t.rows:
            row_data = [c.text.strip() for c in row.cells]
            if any(cell != "" for cell in row_data):
                data.append(row_data)
        if data:
            tables.append(NormalizedTable(page=1, data=data))

    # Images (best-effort)
    images: List[NormalizedImage] = []
    try:
        seen_hashes: set[int] = set()
        for rel in doc.part.rels.values():
            if "image" not in str(getattr(rel, "reltype", "")):
                continue
            part = getattr(rel, "target_part", None)
            blob = getattr(part, "blob", None)
            if not isinstance(blob, (bytes, bytearray)) or not blob:
                continue
            h = hash(bytes(blob))
            if h in seen_hashes:
                continue
            seen_hashes.add(h)
            images.append(
                NormalizedImage(
                    page=1,
                    base64=base64.b64encode(bytes(blob)).decode("ascii"),
                    caption="",
                )
            )
    except Exception:
        images = []

    # Structure sections: headings (best-effort)
    sections: List[Dict[str, Any]] = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style_name = str(getattr(getattr(p, "style", None), "name", ""))
        if style_name.startswith("Heading"):
            sections.append({"kind": "heading", "page": 1, "title": txt})

    return {
        "text_content": text_content,
        "tables": tables,
        "images": images,
        "pages": 1,
        "sections": sections,
    }


def _parse_csv(raw_csv: bytes) -> Dict[str, Any]:
    # Try utf-8, then fall back.
    try:
        text = raw_csv.decode("utf-8-sig")
    except Exception:
        text = raw_csv.decode("latin-1")

    reader = csv.reader(text.splitlines())
    data: List[List[str]] = []
    for row in reader:
        row_data = [str(c).strip() for c in row]
        if any(c != "" for c in row_data):
            data.append(row_data)

    tables: List[NormalizedTable] = []
    if data:
        tables.append(NormalizedTable(page=1, data=data))

    return {
        "text_content": text.strip(),
        "tables": tables,
        "images": [],
        "pages": 1,
        "sections": [{"kind": "csv", "page": 1, "rows": len(data)}],
    }


async def _ensure_table(pool: Any, *, table: str) -> None:
    create_sql = f"""
    CREATE TABLE IF NOT EXISTS {table} (
      document_id UUID PRIMARY KEY,
      source_url TEXT NOT NULL,
      mime_type TEXT NOT NULL,
      metadata JSONB NOT NULL,
      normalized_json JSONB NOT NULL,
      created_at TIMESTAMPTZ NOT NULL
    );
    """
    async with pool.acquire() as conn:
        await conn.execute(create_sql)


async def _store_normalized(
    pool: Any, *, table: str, ingested: IngestedDocument, normalized: NormalizedDocument
) -> None:
    insert_sql = f"""
    INSERT INTO {table} (document_id, source_url, mime_type, metadata, normalized_json, created_at)
    VALUES ($1, $2, $3, $4::jsonb, $5::jsonb, $6)
    ON CONFLICT (document_id) DO UPDATE SET
      source_url = EXCLUDED.source_url,
      mime_type = EXCLUDED.mime_type,
      metadata = EXCLUDED.metadata,
      normalized_json = EXCLUDED.normalized_json,
      created_at = EXCLUDED.created_at;
    """

    normalized_json = json.dumps(normalized.model_dump(mode="json"))
    metadata_json = json.dumps(ingested.metadata)

    async with pool.acquire() as conn:
        await conn.execute(
            insert_sql,
            ingested.document_id,
            ingested.source_url,
            ingested.mime_type,
            metadata_json,
            normalized_json,
            _utc_now_iso(),
        )


async def _init_pg_pool_with_retry(*, dsn: str, table: str) -> Any:
    import asyncpg

    max_attempts = int(os.getenv("NORMALIZER_PG_CONNECT_MAX_ATTEMPTS", "15"))
    delay = float(os.getenv("NORMALIZER_PG_CONNECT_INITIAL_DELAY_SECONDS", "0.5"))
    delay_max = float(os.getenv("NORMALIZER_PG_CONNECT_MAX_DELAY_SECONDS", "8"))

    last_exc: Exception | None = None
    for attempt in range(1, max_attempts + 1):
        try:
            pool = await asyncpg.create_pool(dsn=dsn, min_size=1, max_size=5)
            await _ensure_table(pool, table=table)
            return pool
        except Exception as e:
            last_exc = e
            if attempt >= max_attempts:
                break
            await asyncio.sleep(delay)
            delay = min(delay * 2, delay_max)

    raise RuntimeError(f"Failed to initialize PostgreSQL pool after {max_attempts} attempts: {last_exc}")


@asynccontextmanager
async def _lifespan(app: FastAPI):
    app.state.pg_pool = None
    app.state.pg_table = _postgres_table()

    dsn = _postgres_dsn()
    if dsn:
        app.state.pg_pool = await _init_pg_pool_with_retry(dsn=dsn, table=app.state.pg_table)

    try:
        yield
    finally:
        pool = getattr(app.state, "pg_pool", None)
        if pool is not None:
            await pool.close()


def create_app() -> FastAPI:
    app = FastAPI(title="format-normalizer", version="1.0.0", lifespan=_lifespan)

    @app.get("/healthz")
    async def healthz() -> Dict[str, str]:
        return {"status": "ok"}

    @app.post("/normalize", response_model=NormalizedDocument)
    async def normalize(doc: IngestedDocument) -> NormalizedDocument:
        raw = _b64_decode(doc.raw_bytes)
        kind = _detect_kind(ingested=doc, raw=raw)

        if kind == "pdf":
            parsed = parse_pdf(raw)
        elif kind == "xlsx":
            parsed = parse_xlsx(raw)
        elif kind == "docx":
            parsed = _parse_docx(raw)
        elif kind == "csv":
            parsed = _parse_csv(raw)
        else:
            raise HTTPException(
                status_code=415,
                detail={
                    "error": "unsupported_document_type",
                    "message": "Supported types: PDF, DOCX, XLSX, CSV.",
                    "mime_type": doc.mime_type,
                },
            )

        structure = NormalizedStructure(pages=int(parsed["pages"]), sections=list(parsed.get("sections", [])))

        normalized = NormalizedDocument(
            document_id=doc.document_id,
            text_content=str(parsed.get("text_content", "")),
            tables=list(parsed.get("tables", [])),
            images=list(parsed.get("images", [])),
            structure=structure,
        )

        pool = getattr(app.state, "pg_pool", None)
        table = getattr(app.state, "pg_table", "normalized_documents")
        if pool is not None:
            try:
                await _store_normalized(pool, table=table, ingested=doc, normalized=normalized)
            except Exception as e:
                raise HTTPException(
                    status_code=502,
                    detail={
                        "error": "postgres_write_failed",
                        "message": str(e),
                    },
                ) from e

        return normalized

    return app


app = create_app()


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8001")))
